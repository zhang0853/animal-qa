"""
动物检疫知识库核心类 - 支持只读模式
"""

import os
import sys
import hashlib
import pickle
import json
import threading
import requests
import warnings
import logging
import random
from typing import List, Dict, Any, Optional, Generator
from datetime import datetime

# ========== 添加Word导出相关库 ==========
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ======================================

# ========== 抑制各种警告 ==========
os.environ["QT_LOGGING_RULES"] = "*.debug=false;qt.*.warning=false"
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# 屏蔽pypdf的特定警告
warnings.filterwarnings("ignore", message="Multiple definitions in dictionary")
warnings.filterwarnings("ignore", message="EOF marker")
warnings.filterwarnings("ignore", category=UserWarning, module="pypdf")
logging.getLogger("pypdf").setLevel(logging.ERROR)
# ==================================

# ==================== LangChain相关导入 ====================
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI

# ==================== 配置参数 ====================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
KNOWLEDGE_BASE_DIR = os.path.join(BASE_DIR, "检疫知识库")
INDEX_SAVE_PATH = os.path.join(BASE_DIR, "faiss_data")
CACHE_FILE = os.path.join(BASE_DIR, "processed.cache")
CONFIG_FILE = os.path.join(BASE_DIR, "config.json")
EXPORT_DIR = os.path.join(BASE_DIR, "导出问答")

# 只读模式配置（通过环境变量控制）
READONLY_MODE = os.environ.get("KB_READONLY", "true").lower() == "true"
print(f"📚 知识库模式: {'只读' if READONLY_MODE else '可写'}")

# 博查API配置
BOCHA_API_KEY = "sk-cbb1f9c3978c44e095a3f5744a39c277"
BOCHA_API_URL = "https://api.bochaai.com/v1/web-search"

# DeepSeek API配置
DEEPSEEK_API_KEY = "sk-3f04135d9ffb4636985c3796b6562df0"
DEEPSEEK_API_BASE = "https://api.deepseek.com/v1"

# 支持的文件格式
LOADER_MAPPING = {
    '.pdf': PyPDFLoader,
    '.docx': Docx2txtLoader,
    '.doc': Docx2txtLoader,
    '.txt': TextLoader,
}


# ==================== 配置管理 ====================
def load_config():
    """加载配置文件"""
    default_config = {
        "local_results_count": 5,
        "deepseek_model": "deepseek-chat",
        "temperature": 0.8,
        "max_tokens": 2000,
        "chunk_size": 500,
        "chunk_overlap": 100,
        "auto_rebuild": True,
        "enable_web_search": True,
        "web_search_count": 3,
        "web_search_freshness": "noLimit",
        "theme": "light",
        "font_size": 11,
        "max_history_turns": 10,
        "export_author": "动物检疫专家"
    }

    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return {**default_config, **config}
        except:
            return default_config
    return default_config


def save_config(config):
    """保存配置文件"""
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


# ==================== 博查联网搜索 ====================
def bocha_web_search(query: str, count: int = 3, freshness: str = "noLimit") -> List[Dict]:
    """调用博查Web Search API"""
    if not BOCHA_API_KEY or BOCHA_API_KEY == "你的博查API密钥":
        return []

    headers = {
        "Authorization": f"Bearer {BOCHA_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "query": query,
        "summary": True,
        "count": min(count, 50),
        "freshness": freshness
    }

    try:
        response = requests.post(BOCHA_API_URL, headers=headers, json=payload, timeout=10)
        if response.status_code == 200:
            data = response.json()
            web_pages = data.get("data", {}).get("webPages", {}).get("value", [])

            results = []
            for page in web_pages:
                results.append({
                    "title": page.get("name", ""),
                    "url": page.get("url", ""),
                    "snippet": page.get("snippet", ""),
                    "site_name": page.get("siteName", ""),
                    "date": page.get("dateLastCrawled", ""),
                    "type": "web"
                })
            return results
        else:
            return []
    except Exception as e:
        print(f"博查搜索异常: {e}")
        return []


# ==================== Word导出功能 ====================
def set_run_font(run, size=None, bold=False, color=None, name='宋体'):
    """统一设置Run的字体属性"""
    if size:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def set_paragraph_font(paragraph, size=None, name='宋体'):
    """设置段落默认字体"""
    style = paragraph.style
    style.font.name = name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size:
        style.font.size = Pt(size)


def export_single_qa(question: str, answer: str, sources: List[Dict], export_path: str):
    """导出单个问答到Word文档"""
    doc = Document()

    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    set_paragraph_font(header_para, size=11)
    header_run = header_para.add_run("动物检疫知识库智能问答系统")
    set_run_font(header_run, size=11, color=RGBColor(128, 128, 128))
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("动物检疫知识库问答记录")
    set_run_font(title_run, size=20, bold=True, name='黑体')

    doc.add_paragraph()

    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5)

    cell_time_label = table.cell(0, 0)
    cell_time_label.text = "生成时间："
    set_paragraph_font(cell_time_label.paragraphs[0])
    cell_time_label.paragraphs[0].runs[0].bold = True

    cell_time_value = table.cell(0, 1)
    cell_time_value.text = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
    set_paragraph_font(cell_time_value.paragraphs[0])

    cell_author_label = table.cell(1, 0)
    cell_author_label.text = "导出作者："
    set_paragraph_font(cell_author_label.paragraphs[0])
    cell_author_label.paragraphs[0].runs[0].bold = True

    cell_author_value = table.cell(1, 1)
    cell_author_value.text = "动物检疫专家"
    set_paragraph_font(cell_author_value.paragraphs[0])

    doc.add_paragraph()

    question_heading = doc.add_paragraph()
    question_run = question_heading.add_run("❓ 问题")
    set_run_font(question_run, size=14, bold=True, color=RGBColor(26, 115, 232), name='黑体')

    question_table = doc.add_table(rows=1, cols=1)
    question_table.autofit = False
    question_table.columns[0].width = Inches(6.5)
    question_cell = question_table.cell(0, 0)
    question_cell.text = question
    set_paragraph_font(question_cell.paragraphs[0], size=12)
    question_cell.paragraphs[0].paragraph_format.left_indent = Inches(0.2)
    question_cell.paragraphs[0].paragraph_format.right_indent = Inches(0.2)

    tc = question_cell._tc
    tcPr = tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), 'F5F5F5')
    tcPr.append(sh)

    doc.add_paragraph()

    answer_heading = doc.add_paragraph()
    answer_run = answer_heading.add_run("💡 答案")
    set_run_font(answer_run, size=14, bold=True, color=RGBColor(26, 115, 232), name='黑体')

    answer_table = doc.add_table(rows=1, cols=1)
    answer_table.autofit = False
    answer_table.columns[0].width = Inches(6.5)
    answer_cell = answer_table.cell(0, 0)

    paragraphs = answer.split('\n')
    for i, para_text in enumerate(paragraphs):
        if para_text.strip():
            p = answer_cell.paragraphs[0] if i == 0 else answer_cell.add_paragraph()
            p.text = para_text
            set_paragraph_font(p, size=12)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.line_spacing = 1.5

    tc = answer_cell._tc
    tcPr = tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), 'E8F0FE')
    tcPr.append(sh)

    doc.add_paragraph()

    if sources:
        sources_heading = doc.add_paragraph()
        sources_run = sources_heading.add_run("📚 信息来源")
        set_run_font(sources_run, size=14, bold=True, color=RGBColor(26, 115, 232), name='黑体')

        local_sources = [s for s in sources if s.get('type') == 'local']
        web_sources = [s for s in sources if s.get('type') == 'web']

        if local_sources:
            local_sub = doc.add_paragraph()
            local_sub_run = local_sub.add_run("本地知识库")
            set_run_font(local_sub_run, size=12, bold=True, color=RGBColor(0, 100, 0))

            local_table = doc.add_table(rows=len(local_sources), cols=2)
            local_table.style = 'Light Shading Accent 3'
            local_table.autofit = False
            local_table.columns[0].width = Inches(1)
            local_table.columns[1].width = Inches(5.5)

            for i, source in enumerate(local_sources):
                cell_num = local_table.cell(i, 0)
                cell_num.text = str(i + 1)
                set_paragraph_font(cell_num.paragraphs[0])
                cell_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_content = local_table.cell(i, 1)
                content_text = f"{source['source']}\n{source['content'][:200]}..."
                cell_content.text = content_text
                set_paragraph_font(cell_content.paragraphs[0], size=10)

            doc.add_paragraph()

        if web_sources:
            web_sub = doc.add_paragraph()
            web_sub_run = web_sub.add_run("网络信息")
            set_run_font(web_sub_run, size=12, bold=True, color=RGBColor(200, 100, 0))

            web_table = doc.add_table(rows=len(web_sources), cols=2)
            web_table.style = 'Light Shading Accent 2'
            web_table.autofit = False
            web_table.columns[0].width = Inches(1)
            web_table.columns[1].width = Inches(5.5)

            for i, source in enumerate(web_sources):
                cell_num = web_table.cell(i, 0)
                cell_num.text = str(i + 1)
                set_paragraph_font(cell_num.paragraphs[0])
                cell_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_content = web_table.cell(i, 1)
                content_text = f"标题：{source.get('title', '无标题')}\n来源：{source['source']}\n摘要：{source['content'][:200]}..."
                cell_content.text = content_text
                set_paragraph_font(cell_content.paragraphs[0], size=10)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_text = f"第 1 页 | 动物检疫知识库智能问答系统 | {datetime.now().strftime('%Y年%m月%d日')}"
    footer_run = footer_para.add_run(footer_text)
    set_run_font(footer_run, size=9, color=RGBColor(128, 128, 128))
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(export_path)


def export_batch_qa(qa_list: List[Dict], export_path: str):
    """批量导出问答到Word文档"""
    doc = Document()

    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_run = header_para.add_run("动物检疫知识库智能问答系统 - 批量导出")
    set_run_font(header_run, size=11, color=RGBColor(128, 128, 128))
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("动物检疫知识库问答记录（批量导出）")
    set_run_font(title_run, size=20, bold=True, name='黑体')

    doc.add_paragraph()

    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = 'Table Grid'
    info_table.autofit = False
    info_table.columns[0].width = Inches(1.5)
    info_table.columns[1].width = Inches(5)

    cell_time_label = info_table.cell(0, 0)
    cell_time_label.text = "导出时间："
    set_paragraph_font(cell_time_label.paragraphs[0])
    cell_time_label.paragraphs[0].runs[0].bold = True

    cell_time_value = info_table.cell(0, 1)
    cell_time_value.text = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
    set_paragraph_font(cell_time_value.paragraphs[0])

    cell_count_label = info_table.cell(1, 0)
    cell_count_label.text = "导出数量："
    set_paragraph_font(cell_count_label.paragraphs[0])
    cell_count_label.paragraphs[0].runs[0].bold = True

    cell_count_value = info_table.cell(1, 1)
    cell_count_value.text = f"{len(qa_list)} 条记录"
    set_paragraph_font(cell_count_value.paragraphs[0])
    cell_count_value.paragraphs[0].runs[0].bold = True
    cell_count_value.paragraphs[0].runs[0].font.color.rgb = RGBColor(26, 115, 232)

    doc.add_paragraph()

    toc_heading = doc.add_paragraph()
    toc_run = toc_heading.add_run("📑 目录")
    set_run_font(toc_run, size=16, bold=True, color=RGBColor(26, 115, 232), name='黑体')

    toc_table = doc.add_table(rows=len(qa_list) + 1, cols=2)
    toc_table.style = 'Light List Accent 1'
    toc_table.autofit = False
    toc_table.columns[0].width = Inches(1)
    toc_table.columns[1].width = Inches(6)

    header_num = toc_table.cell(0, 0)
    header_num.text = "序号"
    set_paragraph_font(header_num.paragraphs[0])
    header_num.paragraphs[0].runs[0].bold = True
    header_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    header_title = toc_table.cell(0, 1)
    header_title.text = "问题"
    set_paragraph_font(header_title.paragraphs[0])
    header_title.paragraphs[0].runs[0].bold = True
    header_title.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, qa in enumerate(qa_list, 1):
        cell_num = toc_table.cell(i, 0)
        cell_num.text = str(i)
        set_paragraph_font(cell_num.paragraphs[0])
        cell_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_question = toc_table.cell(i, 1)
        question_preview = qa['question'][:50] + "..." if len(qa['question']) > 50 else qa['question']
        cell_question.text = question_preview
        set_paragraph_font(cell_question.paragraphs[0])

    doc.add_page_break()

    for idx, qa in enumerate(qa_list, 1):
        number_heading = doc.add_paragraph()
        number_run = number_heading.add_run(f"第 {idx} 条 ｜ {qa.get('timestamp', '')}")
        set_run_font(number_run, size=14, bold=True, color=RGBColor(100, 100, 100), name='黑体')

        question_heading = doc.add_paragraph()
        question_run = question_heading.add_run("❓ 问题")
        set_run_font(question_run, size=13, bold=True, color=RGBColor(26, 115, 232))

        question_table = doc.add_table(rows=1, cols=1)
        question_table.autofit = False
        question_table.columns[0].width = Inches(6.5)
        question_cell = question_table.cell(0, 0)
        question_cell.text = qa['question']
        set_paragraph_font(question_cell.paragraphs[0], size=11)

        tc = question_cell._tc
        tcPr = tc.get_or_add_tcPr()
        sh = OxmlElement('w:shd')
        sh.set(qn('w:fill'), 'F5F5F5')
        tcPr.append(sh)

        doc.add_paragraph()

        answer_heading = doc.add_paragraph()
        answer_run = answer_heading.add_run("💡 答案")
        set_run_font(answer_run, size=13, bold=True, color=RGBColor(26, 115, 232))

        answer_table = doc.add_table(rows=1, cols=1)
        answer_table.autofit = False
        answer_table.columns[0].width = Inches(6.5)
        answer_cell = answer_table.cell(0, 0)

        paragraphs = qa['answer'].split('\n')
        for i, para_text in enumerate(paragraphs):
            if para_text.strip():
                p = answer_cell.paragraphs[0] if i == 0 else answer_cell.add_paragraph()
                p.text = para_text
                set_paragraph_font(p, size=11)
                p.paragraph_format.line_spacing = 1.3

        tc = answer_cell._tc
        tcPr = tc.get_or_add_tcPr()
        sh = OxmlElement('w:shd')
        sh.set(qn('w:fill'), 'E8F0FE')
        tcPr.append(sh)

        if qa.get('sources'):
            doc.add_paragraph()
            sources_heading = doc.add_paragraph()
            sources_run = sources_heading.add_run("📚 信息来源")
            set_run_font(sources_run, size=12, bold=True, color=RGBColor(26, 115, 232))

            for source in qa['sources'][:3]:
                p = doc.add_paragraph(style='List Bullet')
                if source.get('type') == 'local':
                    p.add_run(f"本地：{source['source']}")
                else:
                    p.add_run(f"网络：{source.get('title', '无标题')}")
                set_paragraph_font(p, size=10)

        if idx < len(qa_list):
            doc.add_page_break()

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_text = f"动物检疫知识库智能问答系统 | 共 {len(qa_list)} 条记录 | {datetime.now().strftime('%Y年%m月%d日')}"
    footer_run = footer_para.add_run(footer_text)
    set_run_font(footer_run, size=9, color=RGBColor(128, 128, 128))
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(export_path)


# ==================== 历史记录项类 ====================
class HistoryItem:
    """历史记录项，存储完整的问答数据"""

    def __init__(self, question: str, answer: str, sources: List[Dict], timestamp: datetime = None):
        self.question = question
        self.answer = answer
        self.sources = sources
        self.timestamp = timestamp or datetime.now()
        self.selected = False

    def to_dict(self):
        return {
            "role": "user",
            "content": self.question
        }, {
            "role": "assistant",
            "content": self.answer
        }

    def to_export_dict(self):
        return {
            "question": self.question,
            "answer": self.answer,
            "sources": self.sources,
            "timestamp": self.timestamp.strftime("%Y-%m-%d %H:%M:%S")
        }


# ==================== 知识库核心类（支持只读模式）====================
class KnowledgeBase:
    """知识库管理类"""

    def __init__(self):
        self.config = load_config()
        self.vectorstore = None
        self.embeddings = None
        self.retriever = None
        self.llm = None
        self.readonly = READONLY_MODE

        self._init_embeddings_simple()
        self._init_llm()

    def _init_embeddings_simple(self):
        """简化版Embedding初始化"""
        try:
            import torch

            if torch.cuda.is_available():
                device = "cuda"
                gpu_name = torch.cuda.get_device_name(0)
                print(f"✅ 检测到GPU：{gpu_name}，启用GPU加速")
            else:
                device = "cpu"
                print("⚠️ GPU不可用，使用CPU")

            self.embeddings = HuggingFaceEmbeddings(
                model_name="BAAI/bge-small-zh-v1.5",
                model_kwargs={'device': device},
                encode_kwargs={'normalize_embeddings': True}
            )

            test_embed = self.embeddings.embed_query("测试")
            print(f"✅ Embedding模型初始化成功，向量维度：{len(test_embed)}")

        except Exception as e:
            print(f"❌ Embedding初始化失败: {e}")
            try:
                self.embeddings = HuggingFaceEmbeddings(
                    model_name="BAAI/bge-small-zh-v1.5",
                    model_kwargs={'device': 'cpu'},
                    encode_kwargs={'normalize_embeddings': True}
                )
                print("✅ CPU模式初始化成功")
            except Exception as e2:
                print(f"❌ CPU模式也失败: {e2}")
                self.embeddings = None

    def _init_llm(self):
        """初始化DeepSeek（高多样性模式）"""
        try:
            self.llm = ChatOpenAI(
                model=self.config.get("deepseek_model", "deepseek-chat"),
                openai_api_key=DEEPSEEK_API_KEY,
                openai_api_base=DEEPSEEK_API_BASE,
                temperature=self.config.get("temperature", 0.8),
                max_tokens=self.config.get("max_tokens", 2000),
                streaming=True,
                top_p=0.9,
                frequency_penalty=0.3,
                presence_penalty=0.3
            )
            # 测试连接
            self.llm.invoke("test")
            print("✅ DeepSeek连接成功（高多样性模式）")
        except Exception as e:
            print(f"❌ DeepSeek连接失败: {e}")
            self.llm = None

    def is_readonly(self):
        """返回当前是否为只读模式"""
        return self.readonly

    def should_update_index(self):
        """判断是否需要更新索引（只读模式下永远返回False）"""
        import os
        if os.environ.get("CLOUD_MODE", "false") == "true":
           print("☁️ 云端模式，跳过索引更新检查")
           return False
        
        
        if self.readonly:
            print("⚠️ 只读模式，跳过索引更新检查")
            return False

        if not self.embeddings:
            print("❌ Embedding模型未初始化")
            return False

        index_file = os.path.join(INDEX_SAVE_PATH, "index.faiss")
        if not os.path.exists(index_file):
            print("索引文件不存在，需要重建")
            return True

        if not os.path.exists(CACHE_FILE):
            print("缓存文件不存在，需要重建")
            return True

        if self.config.get("auto_rebuild", True):
            current_count = 0
            for root, dirs, files in os.walk(KNOWLEDGE_BASE_DIR):
                for file in files:
                    ext = os.path.splitext(file)[1].lower()
                    if ext in LOADER_MAPPING:
                        current_count += 1

            try:
                with open(CACHE_FILE, 'rb') as f:
                    processed = pickle.load(f)
                    cached_count = len(processed)
                    if current_count != cached_count:
                        print(f"文件数量变化：{cached_count} -> {current_count}，需要重建")
                        return True
            except:
                return True

        return False

    def rebuild_index(self, progress_callback=None):
        """重建索引（只读模式下禁止）"""
        if self.readonly:
            error_msg = "❌ 当前为只读模式，无法重建索引。请使用管理员后台进行更新。"
            print(error_msg)
            return False, error_msg

        if not self.embeddings:
            return False, "Embedding模型初始化失败，无法重建索引"

        if progress_callback:
            progress_callback(0, "开始重建索引...")

        if not os.path.exists(KNOWLEDGE_BASE_DIR):
            os.makedirs(KNOWLEDGE_BASE_DIR)
            return False, f"知识库目录不存在，已创建：{KNOWLEDGE_BASE_DIR}，请放入文档后重试"

        documents = self._load_all_documents(progress_callback)
        if not documents:
            return False, f"没有成功加载任何文档，请检查 {KNOWLEDGE_BASE_DIR} 目录下的文件格式（支持PDF、DOCX、TXT）"

        if progress_callback:
            progress_callback(30, f"已加载 {len(documents)} 个文档片段")

        chunks = self._split_documents(documents)
        if not chunks:
            return False, "切分文档失败"

        if progress_callback:
            progress_callback(60, f"切分成 {len(chunks)} 个片段")

        if progress_callback:
            progress_callback(70, "正在构建向量索引...")

        os.makedirs(INDEX_SAVE_PATH, exist_ok=True)

        texts = [chunk.page_content for chunk in chunks]
        metadatas = [chunk.metadata for chunk in chunks]

        try:
            # 分批处理，避免内存问题
            batch_size = 1000
            if len(texts) > batch_size:
                print(f"文档数量较大({len(texts)})，分批处理...")
                self.vectorstore = FAISS.from_texts(
                    texts=texts[:batch_size],
                    embedding=self.embeddings,
                    metadatas=metadatas[:batch_size]
                )
                for i in range(batch_size, len(texts), batch_size):
                    end = min(i + batch_size, len(texts))
                    print(f"添加批次 {i // batch_size + 1}/{(len(texts) - 1) // batch_size + 1}")
                    self.vectorstore.add_texts(
                        texts=texts[i:end],
                        metadatas=metadatas[i:end]
                    )
            else:
                self.vectorstore = FAISS.from_texts(
                    texts=texts,
                    embedding=self.embeddings,
                    metadatas=metadatas
                )

            self.vectorstore.save_local(INDEX_SAVE_PATH)
            self.retriever = self.vectorstore.as_retriever(
                search_kwargs={"k": self.config.get("local_results_count", 5)}
            )

            if progress_callback:
                progress_callback(100, f"索引创建成功，包含 {len(texts)} 个片段")

            print(f"✅ 索引重建成功，共 {len(texts)} 个片段")
            return True, f"索引创建成功，包含 {len(texts)} 个片段"
        except Exception as e:
            print(f"❌ 索引创建失败: {str(e)}")
            import traceback
            traceback.print_exc()
            return False, f"索引创建失败: {str(e)}"

    def load_index(self):
        """加载已有索引"""
        if not self.embeddings:
            return False, "Embedding模型初始化失败，无法加载索引"

        try:
            self.vectorstore = FAISS.load_local(
                INDEX_SAVE_PATH,
                self.embeddings,
                allow_dangerous_deserialization=True
            )
            self.retriever = self.vectorstore.as_retriever(
                search_kwargs={"k": self.config.get("local_results_count", 5)}
            )

            if hasattr(self.vectorstore, 'index') and self.vectorstore.index.ntotal > 0:
                print(f"✅ 索引加载成功，共 {self.vectorstore.index.ntotal} 个向量")
                return True, f"索引加载成功，包含 {self.vectorstore.index.ntotal} 个片段"
            else:
                print("⚠️ 索引文件为空，需要重建")
                return False, "索引文件为空，需要重建"
        except Exception as e:
            print(f"❌ 索引加载失败: {str(e)}")
            return False, f"索引加载失败: {str(e)}"

    def _load_all_documents(self, progress_callback=None):
        """加载所有文档"""
        import warnings
        from pypdf.errors import PdfReadWarning

        warnings.filterwarnings("ignore", category=PdfReadWarning)

        all_files = []
        for root, dirs, files in os.walk(KNOWLEDGE_BASE_DIR):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in LOADER_MAPPING:
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, KNOWLEDGE_BASE_DIR)
                    all_files.append((file_path, ext, rel_path))
                    print(f"找到文件：{rel_path}")

        if not all_files:
            print("⚠️ 知识库目录中没有找到支持的文档")
            return []

        documents = []
        for i, (file_path, ext, rel_path) in enumerate(all_files):
            if progress_callback:
                progress = int(20 * (i + 1) / len(all_files))
                progress_callback(progress, f"加载文件 {i + 1}/{len(all_files)}: {rel_path}")

            try:
                print(f"正在加载：{rel_path}")
                if ext == '.pdf':
                    loader = PyPDFLoader(file_path, extract_images=False)
                    docs = loader.load()
                    if len(docs) > 100:
                        docs = docs[:100]
                else:
                    loader = LOADER_MAPPING[ext](file_path)
                    docs = loader.load()

                for doc in docs:
                    doc.metadata["source"] = rel_path
                    doc.metadata["file_name"] = os.path.basename(file_path)
                documents.extend(docs)
                print(f"✅ 加载成功：{rel_path} ({len(docs)}页/段)")
            except Exception as e:
                print(f"❌ 加载失败 {rel_path}: {e}")
                continue

        processed = {}
        for file_path, ext, rel_path in all_files:
            try:
                with open(file_path, 'rb') as f:
                    processed[file_path] = hashlib.md5(f.read()).hexdigest()
            except:
                pass

        with open(CACHE_FILE, 'wb') as f:
            pickle.dump(processed, f)

        print(f"✅ 共加载 {len(documents)} 个文档片段")
        return documents

    def _split_documents(self, documents):
        """切分文档"""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.get("chunk_size", 500),
            chunk_overlap=self.config.get("chunk_overlap", 100),
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
        )
        return splitter.split_documents(documents)

    def search_local(self, query: str) -> List[Dict]:
        """本地搜索（随机化检索结果）"""
        if not self.retriever:
            print("❌ 检索器未初始化")
            return []

        try:
            docs_with_scores = self.vectorstore.similarity_search_with_score(query, k=20)
            print(f"🔍 本地检索到 {len(docs_with_scores)} 个候选片段")

            weighted_docs = []
            for doc, score in docs_with_scores:
                similarity = 1 - score if score <= 1 else 1 / (1 + score)
                if similarity > 0.2:
                    weight = similarity * (0.8 + 0.4 * random.random())
                    weighted_docs.append((doc, weight, similarity))

            weighted_docs.sort(key=lambda x: x[1], reverse=True)

            k = self.config.get("local_results_count", 5)
            selected = weighted_docs[:k]

            results = []
            for doc, weight, similarity in selected:
                results.append({
                    "content": doc.page_content,
                    "source": doc.metadata.get("file_name", "未知"),
                    "file_path": doc.metadata.get("source", "未知"),
                    "similarity": f"{similarity:.2f}",
                    "type": "local"
                })

            print(f"  随机加权选择 {len(results)} 个片段")
            for i, r in enumerate(results):
                print(f"  结果{i + 1}: {r['source']} - {r['content'][:50]}...")

            return results

        except Exception as e:
            print(f"❌ 本地检索失败: {e}")
            return []

    def search_web(self, query: str) -> List[Dict]:
        """联网搜索"""
        if not self.config.get("enable_web_search", True):
            return []

        count = self.config.get("web_search_count", 3)
        freshness = self.config.get("web_search_freshness", "noLimit")

        web_results = bocha_web_search(query, count, freshness)
        print(f"🌐 联网检索到 {len(web_results)} 条结果")

        results = []
        for item in web_results:
            results.append({
                "content": f"标题：{item['title']}\n摘要：{item['snippet']}",
                "source": item['url'],
                "title": item['title'],
                "site": item['site_name'],
                "date": item['date'],
                "type": "web"
            })

        return results

    def _get_answer_style(self, question: str) -> str:
        """根据问题类型选择回答风格"""
        styles = [
            "请用通俗易懂的语言解释",
            "请从专业角度详细说明",
            "请用步骤化的方式说明",
            "请结合实际应用场景说明",
            "请用对比的方式说明"
        ]

        if any(word in question for word in ["怎么", "如何", "步骤"]):
            return "请用步骤化的方式说明，可以分点列出"
        elif any(word in question for word in ["区别", "不同", "对比"]):
            return "请用对比的方式说明，突出异同点"
        elif any(word in question for word in ["举例", "案例", "比如"]):
            return "请结合实际应用场景说明，可以举一些典型案例"
        else:
            return random.choice(styles)

    def _build_prompt(self, question: str, context: str, history: List[Dict] = None) -> str:
        """构建提示词（多风格轮换）"""
        history_text = ""
        if history and len(history) > 0:
            history_parts = ["以下是之前的对话："]
            for i, turn in enumerate(history[-6:]):
                if isinstance(turn, dict) and 'role' in turn and 'content' in turn:
                    role = "用户" if turn["role"] == "user" else "助手"
                    content = turn['content'][:200] + "..." if len(turn['content']) > 200 else turn['content']
                    history_parts.append(f"{role}：{content}")
                elif isinstance(turn, dict) and 'question' in turn:
                    history_parts.append(f"用户：{turn['question'][:200]}...")
                    if 'answer' in turn:
                        history_parts.append(f"助手：{turn['answer'][:200]}...")
            history_text = "\n".join(history_parts) + "\n\n"

        style = self._get_answer_style(question)

        prompt = f"""{history_text}【参考资料】
{context}

【当前问题】
{question}

【回答要求】
1. {style}
2. 不要直接复制粘贴参考资料，用自己的话重新组织
3. 确保信息准确，不能编造
4. 如果参考资料不足，请说明

请开始回答："""
        return prompt

    def ask_stream(self, question: str, history: List[Dict] = None) -> Generator[Dict, None, None]:
        """
        流式提问
        """
        result = {
            "question": question,
            "answer": "",
            "sources": [],
            "web_used": False
        }

        local_results = self.search_local(question)

        need_web = self.config.get("enable_web_search", True)
        if len(local_results) == 0:
            need_web = True
        elif any(word in question for word in ["最新", "今天", "新闻", "2026", "2025"]):
            need_web = True

        web_results = []
        if need_web:
            web_results = self.search_web(question)
            result["web_used"] = bool(web_results)

        all_sources = local_results + web_results
        result["sources"] = all_sources

        if not all_sources:
            result["answer"] = "未找到相关信息"
            yield {"type": "complete", "answer": result["answer"], "sources": [], "web_used": False}
            return

        if self.llm:
            try:
                context_parts = []

                if local_results:
                    context_parts.append("【本地知识库】")
                    for doc in local_results:
                        context_parts.append(doc['content'])

                if web_results:
                    context_parts.append("【联网搜索结果】")
                    for doc in web_results:
                        context_parts.append(f"标题：{doc['title']}\n摘要：{doc['content']}")

                context = "\n\n".join(context_parts)

                prompt = self._build_prompt(question, context, history)

                full_answer = ""
                for chunk in self.llm.stream(prompt):
                    if hasattr(chunk, 'content'):
                        content = chunk.content
                    else:
                        content = str(chunk)

                    full_answer += content
                    yield {
                        "type": "chunk",
                        "content": content,
                        "full_answer": full_answer
                    }

                result["answer"] = full_answer

                yield {"type": "complete", "answer": full_answer, "sources": all_sources,
                       "web_used": result["web_used"]}

            except Exception as e:
                error_msg = f"生成答案时出错: {str(e)}"
                print(f"❌ {error_msg}")
                import traceback
                traceback.print_exc()
                yield {"type": "error", "content": error_msg}
        else:
            fallback_answer = self._format_sources(all_sources)
            yield {"type": "complete", "answer": fallback_answer, "sources": all_sources,
                   "web_used": result["web_used"]}

    def _format_sources(self, sources: List[Dict]) -> str:
        """格式化来源"""
        if not sources:
            return "未找到相关信息"

        output = []

        local = [s for s in sources if s.get("type") == "local"]
        web = [s for s in sources if s.get("type") == "web"]

        if local:
            output.append("📚 本地资料：")
            for i, doc in enumerate(local, 1):
                output.append(f"{i}. {doc['source']}")
                output.append(f"   {doc['content'][:150]}...")

        if web:
            output.append("\n🌐 网络信息：")
            for i, doc in enumerate(web, 1):
                output.append(f"{i}. {doc.get('title', '无标题')}")
                output.append(f"   {doc['content'][:150]}...")
                output.append(f"   来源：{doc.get('site', doc['source'])}")

        return "\n".join(output)
