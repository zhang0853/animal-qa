"""
动物检疫知识库智能问答系统 - Streamlit网页版
"""

import os
import sys
import hashlib
import pickle
import json
import threading
import requests
import warnings
import logging
import random
from typing import List, Dict, Any, Optional, Generator
from datetime import datetime
import streamlit as st
import pandas as pd
from io import BytesIO

# ========== 添加Word导出相关库 ==========
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========== 抑制各种警告 ==========
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

warnings.filterwarnings("ignore", message="Multiple definitions in dictionary")
warnings.filterwarnings("ignore", message="EOF marker")
warnings.filterwarnings("ignore", category=UserWarning, module="pypdf")
logging.getLogger("pypdf").setLevel(logging.ERROR)
# ==================================

# ==================== LangChain相关导入 ====================
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI

# ==================== 配置参数 ====================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
KNOWLEDGE_BASE_DIR = os.path.join(BASE_DIR, "检疫知识库")
INDEX_SAVE_PATH = os.path.join(BASE_DIR, "faiss_data")
CACHE_FILE = os.path.join(BASE_DIR, "processed.cache")
CONFIG_FILE = os.path.join(BASE_DIR, "config.json")
EXPORT_DIR = os.path.join(BASE_DIR, "导出问答")

# API配置
BOCHA_API_KEY = "sk-cbb1f9c3978c44e095a3f5744a39c277"
BOCHA_API_URL = "https://api.bochaai.com/v1/web-search"
DEEPSEEK_API_KEY = "sk-3f04135d9ffb4636985c3796b6562df0"
DEEPSEEK_API_BASE = "https://api.deepseek.com/v1"

# 支持的文件格式
LOADER_MAPPING = {
    '.pdf': PyPDFLoader,
    '.docx': Docx2txtLoader,
    '.doc': Docx2txtLoader,
    '.txt': TextLoader,
}


# ==================== 配置管理 ====================
def load_config():
    """加载配置文件"""
    default_config = {
        "local_results_count": 5,
        "deepseek_model": "deepseek-chat",
        "temperature": 0.8,
        "max_tokens": 2000,
        "chunk_size": 500,
        "chunk_overlap": 100,
        "auto_rebuild": True,
        "enable_web_search": True,
        "web_search_count": 3,
        "web_search_freshness": "noLimit",
        "max_history_turns": 10,
    }

    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return {**default_config, **config}
        except:
            return default_config
    return default_config


def save_config(config):
    """保存配置文件"""
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


# ==================== 博查联网搜索 ====================
def bocha_web_search(query: str, count: int = 3, freshness: str = "noLimit") -> List[Dict]:
    """调用博查Web Search API"""
    if not BOCHA_API_KEY:
        return []

    headers = {
        "Authorization": f"Bearer {BOCHA_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "query": query,
        "summary": True,
        "count": min(count, 50),
        "freshness": freshness
    }

    try:
        response = requests.post(BOCHA_API_URL, headers=headers, json=payload, timeout=10)
        if response.status_code == 200:
            data = response.json()
            web_pages = data.get("data", {}).get("webPages", {}).get("value", [])

            results = []
            for page in web_pages:
                results.append({
                    "title": page.get("name", ""),
                    "url": page.get("url", ""),
                    "snippet": page.get("snippet", ""),
                    "site_name": page.get("siteName", ""),
                    "date": page.get("dateLastCrawled", ""),
                    "type": "web"
                })
            return results
        else:
            return []
    except Exception as e:
        print(f"博查搜索异常: {e}")
        return []


# ==================== Word导出功能 ====================
def export_single_qa_to_bytes(question: str, answer: str, sources: List[Dict]) -> BytesIO:
    """导出单个问答到Word文档（返回BytesIO对象）"""
    doc = Document()

    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("动物检疫知识库问答记录")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    # 基本信息
    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = 'Table Grid'

    info_table.cell(0, 0).text = "生成时间："
    info_table.cell(0, 1).text = datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")
    info_table.cell(1, 0).text = "导出作者："
    info_table.cell(1, 1).text = "动物检疫专家"

    doc.add_paragraph()

    # 问题
    q_heading = doc.add_paragraph()
    q_run = q_heading.add_run("❓ 问题")
    q_run.font.size = Pt(14)
    q_run.font.bold = True
    q_run.font.color.rgb = RGBColor(26, 115, 232)

    q_table = doc.add_table(rows=1, cols=1)
    q_table.cell(0, 0).text = question

    doc.add_paragraph()

    # 答案
    a_heading = doc.add_paragraph()
    a_run = a_heading.add_run("💡 答案")
    a_run.font.size = Pt(14)
    a_run.font.bold = True
    a_run.font.color.rgb = RGBColor(26, 115, 232)

    a_table = doc.add_table(rows=1, cols=1)
    a_cell = a_table.cell(0, 0)
    for para in answer.split('\n'):
        if para.strip():
            a_cell.add_paragraph(para)

    # 保存到BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ==================== 知识库核心类 ====================
class KnowledgeBase:
    """知识库管理类"""

    def __init__(self):
        self.config = load_config()
        self.vectorstore = None
        self.embeddings = None
        self.retriever = None
        self.llm = None

        self._init_embeddings_simple()
        self._init_llm()
        self._load_or_build_index()

    def _init_embeddings_simple(self):
        """简化版Embedding初始化"""
        try:
            import torch
            device = "cuda" if torch.cuda.is_available() else "cpu"
            print(f"使用设备: {device}")

            self.embeddings = HuggingFaceEmbeddings(
                model_name="BAAI/bge-small-zh-v1.5",
                model_kwargs={'device': device},
                encode_kwargs={'normalize_embeddings': True}
            )
            print("✅ Embedding模型初始化成功")
        except Exception as e:
            print(f"❌ Embedding初始化失败: {e}")
            self.embeddings = None

    def _init_llm(self):
        """初始化DeepSeek"""
        try:
            self.llm = ChatOpenAI(
                model=self.config.get("deepseek_model", "deepseek-chat"),
                openai_api_key=DEEPSEEK_API_KEY,
                openai_api_base=DEEPSEEK_API_BASE,
                temperature=self.config.get("temperature", 0.8),
                max_tokens=self.config.get("max_tokens", 2000),
                streaming=True,
            )
            print("✅ DeepSeek连接成功")
        except Exception as e:
            print(f"❌ DeepSeek连接失败: {e}")
            self.llm = None

    def _load_or_build_index(self):
        """加载或重建索引"""
        if self.should_update_index():
            self.rebuild_index()
        else:
            self.load_index()

    def should_update_index(self):
        """判断是否需要更新索引"""
        if not self.embeddings:
            return False

        index_file = os.path.join(INDEX_SAVE_PATH, "index.faiss")
        if not os.path.exists(index_file):
            return True

        if not os.path.exists(CACHE_FILE):
            return True

        return False

    def rebuild_index(self, progress_callback=None):
        """重建索引"""
        if not self.embeddings:
            return False, "Embedding模型初始化失败"

        if not os.path.exists(KNOWLEDGE_BASE_DIR):
            os.makedirs(KNOWLEDGE_BASE_DIR)
            return False, f"知识库目录不存在，已创建：{KNOWLEDGE_BASE_DIR}"

        documents = self._load_all_documents()
        if not documents:
            return False, "没有成功加载任何文档"

        chunks = self._split_documents(documents)
        if not chunks:
            return False, "切分文档失败"

        os.makedirs(INDEX_SAVE_PATH, exist_ok=True)

        texts = [chunk.page_content for chunk in chunks]
        metadatas = [chunk.metadata for chunk in chunks]

        try:
            self.vectorstore = FAISS.from_texts(
                texts=texts,
                embedding=self.embeddings,
                metadatas=metadatas
            )
            self.vectorstore.save_local(INDEX_SAVE_PATH)
            self.retriever = self.vectorstore.as_retriever(
                search_kwargs={"k": self.config.get("local_results_count", 5)}
            )
            return True, f"索引创建成功，包含 {len(texts)} 个片段"
        except Exception as e:
            return False, f"索引创建失败: {str(e)}"

    def load_index(self):
        """加载已有索引"""
        if not self.embeddings:
            return False, "Embedding模型初始化失败"

        try:
            self.vectorstore = FAISS.load_local(
                INDEX_SAVE_PATH,
                self.embeddings,
                allow_dangerous_deserialization=True
            )
            self.retriever = self.vectorstore.as_retriever(
                search_kwargs={"k": self.config.get("local_results_count", 5)}
            )
            return True, f"索引加载成功，包含 {self.vectorstore.index.ntotal} 个片段"
        except Exception as e:
            return False, f"索引加载失败: {str(e)}"

    def _load_all_documents(self):
        """加载所有文档"""
        all_files = []
        for root, dirs, files in os.walk(KNOWLEDGE_BASE_DIR):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in LOADER_MAPPING:
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, KNOWLEDGE_BASE_DIR)
                    all_files.append((file_path, ext, rel_path))

        if not all_files:
            return []

        documents = []
        for file_path, ext, rel_path in all_files:
            try:
                if ext == '.pdf':
                    loader = PyPDFLoader(file_path, extract_images=False)
                    docs = loader.load()
                else:
                    loader = LOADER_MAPPING[ext](file_path)
                    docs = loader.load()

                for doc in docs:
                    doc.metadata["source"] = rel_path
                    doc.metadata["file_name"] = os.path.basename(file_path)
                documents.extend(docs)
            except Exception as e:
                print(f"加载失败 {rel_path}: {e}")
                continue

        # 更新缓存
        processed = {}
        for file_path, ext, rel_path in all_files:
            try:
                with open(file_path, 'rb') as f:
                    processed[file_path] = hashlib.md5(f.read()).hexdigest()
            except:
                pass

        with open(CACHE_FILE, 'wb') as f:
            pickle.dump(processed, f)

        return documents

    def _split_documents(self, documents):
        """切分文档"""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.get("chunk_size", 500),
            chunk_overlap=self.config.get("chunk_overlap", 100),
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
        )
        return splitter.split_documents(documents)

    def search_local(self, query: str) -> List[Dict]:
        """本地搜索"""
        if not self.retriever:
            return []

        try:
            docs_with_scores = self.vectorstore.similarity_search_with_score(query, k=20)

            results = []
            for doc, score in docs_with_scores[:self.config.get("local_results_count", 5)]:
                similarity = 1 - score if score <= 1 else 1 / (1 + score)
                results.append({
                    "content": doc.page_content,
                    "source": doc.metadata.get("file_name", "未知"),
                    "file_path": doc.metadata.get("source", "未知"),
                    "similarity": f"{similarity:.2f}",
                    "type": "local"
                })
            return results
        except Exception as e:
            print(f"本地检索失败: {e}")
            return []

    def search_web(self, query: str) -> List[Dict]:
        """联网搜索"""
        if not self.config.get("enable_web_search", True):
            return []

        count = self.config.get("web_search_count", 3)
        freshness = self.config.get("web_search_freshness", "noLimit")
        web_results = bocha_web_search(query, count, freshness)

        results = []
        for item in web_results:
            results.append({
                "content": f"标题：{item['title']}\n摘要：{item['snippet']}",
                "source": item['url'],
                "title": item['title'],
                "site": item['site_name'],
                "type": "web"
            })
        return results

    def ask_stream(self, question: str, history: List[Dict] = None) -> Generator[str, None, None]:
        """流式提问"""
        # 本地搜索
        local_results = self.search_local(question)

        # 联网搜索
        need_web = self.config.get("enable_web_search", True)
        if len(local_results) == 0:
            need_web = True
        elif any(word in question for word in ["最新", "今天", "新闻"]):
            need_web = True

        web_results = self.search_web(question) if need_web else []
        all_sources = local_results + web_results

        if not all_sources:
            yield "未找到相关信息"
            return

        # 生成答案
        if self.llm:
            try:
                # 构建上下文
                context_parts = []
                if local_results:
                    context_parts.append("【本地知识库】")
                    for doc in local_results:
                        context_parts.append(doc['content'])
                if web_results:
                    context_parts.append("【联网搜索结果】")
                    for doc in web_results:
                        context_parts.append(f"标题：{doc['title']}\n摘要：{doc['content']}")
                context = "\n\n".join(context_parts)

                prompt = f"""【参考资料】
{context}

【问题】
{question}

请基于参考资料回答问题，如果资料不足请说明。"""

                for chunk in self.llm.stream(prompt):
                    if hasattr(chunk, 'content'):
                        yield chunk.content
                    else:
                        yield str(chunk)
            except Exception as e:
                yield f"生成答案时出错: {str(e)}"
        else:
            # 降级方案
            result = "【参考资料】\n"
            for src in all_sources[:5]:
                result += f"- {src.get('source', '未知')}: {src['content'][:100]}...\n"
            yield result


# ==================== 初始化知识库 ====================
@st.cache_resource
def init_knowledge_base():
    """初始化知识库（使用Streamlit缓存）"""
    return KnowledgeBase()


# ==================== Streamlit UI ====================
def main():
    st.set_page_config(
        page_title="动物检疫知识库智能问答系统",
        page_icon="🐄",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # 自定义CSS
    st.markdown("""
    <style>
    .stApp {
        background-color: #f5f7fa;
    }
    .main-header {
        background: linear-gradient(135deg, #1a73e8 0%, #1557b0 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        margin-bottom: 2rem;
    }
    .answer-box {
        background-color: #f0f7ff;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #1a73e8;
    }
    .source-item {
        padding: 0.5rem;
        border-bottom: 1px solid #e0e0e0;
    }
    </style>
    """, unsafe_allow_html=True)

    # 标题
    st.markdown("""
    <div class="main-header">
        <h1 style="margin:0; text-align:center;">🐄 动物检疫知识库智能问答系统</h1>
        <p style="text-align:center; margin-top:0.5rem;">基于DeepSeek + 本地知识库 + 联网搜索</p>
    </div>
    """, unsafe_allow_html=True)

    # 初始化知识库
    kb = init_knowledge_base()

    # 初始化会话状态
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "current_answer" not in st.session_state:
        st.session_state.current_answer = ""

    # 侧边栏
    with st.sidebar:
        st.markdown("### ⚙️ 系统设置")

        # 服务状态
        with st.expander("📡 服务状态", expanded=True):
            deepseek_status = "✅ 已连接" if kb.llm else "❌ 未连接"
            st.write(f"DeepSeek: {deepseek_status}")
            st.write(f"博查API: {'✅ 已配置' if BOCHA_API_KEY else '❌ 未配置'}")

            # 知识库统计
            if kb.vectorstore and hasattr(kb.vectorstore, 'index'):
                st.write(f"向量数量: {kb.vectorstore.index.ntotal}")

        # 联网搜索设置
        with st.expander("🌐 联网搜索", expanded=True):
            enable_web = st.checkbox("启用联网搜索", value=kb.config.get("enable_web_search", True))
            web_count = st.slider("返回条数", 1, 10, kb.config.get("web_search_count", 3))
            freshness = st.selectbox(
                "时效性",
                ["noLimit", "oneDay", "oneWeek", "oneMonth", "oneYear"],
                index=0
            )

            if enable_web != kb.config.get("enable_web_search"):
                kb.config["enable_web_search"] = enable_web
                save_config(kb.config)
            if web_count != kb.config.get("web_search_count"):
                kb.config["web_search_count"] = web_count
                save_config(kb.config)

        # 知识库管理
        with st.expander("📚 知识库管理"):
            if st.button("🔄 重建索引", use_container_width=True):
                with st.spinner("正在重建索引..."):
                    success, msg = kb.rebuild_index()
                    if success:
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)

            if st.button("📂 打开知识库文件夹", use_container_width=True):
                os.startfile(KNOWLEDGE_BASE_DIR)

        # 导出功能
        with st.expander("📄 导出功能"):
            if st.button("📥 导出当前对话", use_container_width=True):
                if st.session_state.messages:
                    # 获取最后一条问答
                    last_q = None
                    last_a = None
                    for msg in reversed(st.session_state.messages):
                        if msg["role"] == "assistant":
                            last_a = msg["content"]
                        elif msg["role"] == "user" and last_a:
                            last_q = msg["content"]
                            break

                    if last_q and last_a:
                        doc_bytes = export_single_qa_to_bytes(last_q, last_a, [])
                        st.download_button(
                            label="💾 下载Word文档",
                            data=doc_bytes,
                            file_name=f"问答_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.warning("暂无对话记录")

        # 清空历史
        if st.button("🗑️ 清空对话历史", use_container_width=True):
            st.session_state.messages = []
            st.rerun()

    # 主区域 - 对话历史
    chat_container = st.container()

    with chat_container:
        for message in st.session_state.messages:
            if message["role"] == "user":
                with st.chat_message("user"):
                    st.markdown(message["content"])
            else:
                with st.chat_message("assistant"):
                    st.markdown(message["content"])

    # 输入区域
    with st.container():
        col1, col2 = st.columns([6, 1])
        with col1:
            prompt = st.chat_input("请输入您的问题...")
        with col2:
            if st.button("清空", use_container_width=True):
                st.session_state.messages = []
                st.rerun()

    if prompt:
        # 显示用户问题
        with st.chat_message("user"):
            st.markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})

        # 显示助手回答
        with st.chat_message("assistant"):
            response_placeholder = st.empty()
            full_response = ""

            # 流式输出
            for chunk in kb.ask_stream(prompt, st.session_state.messages):
                full_response += chunk
                response_placeholder.markdown(full_response + "▌")

            response_placeholder.markdown(full_response)

        st.session_state.messages.append({"role": "assistant", "content": full_response})

        # 自动滚动到底部
        st.rerun()


if __name__ == "__main__":
    main()